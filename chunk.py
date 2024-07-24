from langchain.chains.summarize import load_summarize_chain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import PyPDFLoader
from langchain.chat_models import ChatOpenAI
from langchain import PromptTemplate
from docx import Document
import os
from openai import OpenAI
from dotenv import load_dotenv
load_dotenv()
import json


class Chunking:
    api_key = os.getenv("OPENAI_API_KEY")
    os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

    llm=ChatOpenAI(model_name='gpt-3.5-turbo')
    client = OpenAI(api_key = api_key)

    def summarize_cat1(llm,map_prompt_temp,final_combine_prompt,text):

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=100)
        chunks = text_splitter.create_documents([text])
        
        summary_chain = load_summarize_chain(
        llm=llm,
        chain_type='map_reduce',
        map_prompt=map_prompt_temp,
        combine_prompt=final_combine_prompt,
        verbose=False
    )
        output = summary_chain.run(chunks)
        return output
    
    def openai_api(prompt):
        response = Chunking.client.chat.completions.create(
            model = "gpt-3.5-turbo",
            messages=[
                {"role":"system",
                "content":"You are a filenaming system"},
                {"role": "user", "content":prompt}
            ]
        )

        return response.choices[0].message.content

    def invoke(folder_path,sum_folder,random_name):

        chunks_prompt = """
        "You are an AI assistant that consolidates and summarizes information, removing duplicates and maintaining all unique and relevant content. Please consolidate and summarize the following content, removing any duplicate information while keeping all unique and relevant information. You have been given a text document that may contain duplicate lines and redundant information, generate a new text that is unique and devoid of repetitive content. Remove any lines that convey the same meaning or duplicate information.
        Keep in mind the following points:
        - Contains complete information
        - Generate maximum possible content from the given information
        - No repetitive lines
        - No paragraphs or lines that have same information or meaning should occur twice
        - The generated text is unique

        As the following text is the subpart of the full text, so if any of the above mentioned points are there try to solve them accordingly.
        information: {text}
        """
        map_prompt_temp = PromptTemplate(input_variables=['text'],template=chunks_prompt)


        final_combine_prompt ="""
        Your task is to create a new document approximately 3000 words long. The output should maintain the essence of the original document but should eliminate duplicate sentences or sentences conveying the same meaning. Ensure that the new document is coherent, informative, and maintains a consistent style with the original. Deduplicate the entire document and generate a new text based on this text of at least 3000 words. Make sure it follows the following points:

        1. Contains complete information:
        - The new text should have complete information that had been presented in the original text that was provided.
        - No information should be left out or missing from the original text.
        - Entire text should be allowed except lines or paragraphs that have the same meaning that occurs twice.
        - Retain key concepts and ideas while rephrasing or restructuring sentences to avoid redundancy.

        2. Generate maximum possible content from the given information:
        - The new text should be complete and should be as long as possible.
        - The new text should be at least 3000 words.
        - The new text should be completely based on the given text.
        - DO NOT make a summary of the text and generate a bigger text than original.
        - Rephrase if needed but make the input text bigger and not a summary.

        3. No repetitive information:
        - A piece of information once occured should not occur twice. 
        - Make sure the information extracted is unique and does not occur again.
        - Only lines having unique information should be added.
        - Paragraphs having same meaning should not occur again.
        - Single sentences having information that have already occured before should be ignored.
        - The new text should be unique and does not allow repetitive information.

        4. The generated text is unique:
        - The new text that have been extracted from the prompt text should be unique.
        - Every line or paragraph in the new text should be unique and does not contain any repetitive information.
        - Semantic Analysis should be done to the whole text and no lines having similar meaning or information should be repeated twice.

        Data: {text}
        """
        final_combine_prompt_template=PromptTemplate(input_variables=['text'], template=final_combine_prompt)

        # Check if the folder exists
        if os.path.exists(folder_path) and os.path.isdir(folder_path):
            files_list = os.listdir(folder_path)
            if files_list:
                # Store file names in a list
                file_names = []
                for file_name in files_list:
                    file_names.append(file_name)
                
                # Print the list of file names
                print("\nStored file names in a list:")
                print(type(file_names[0]))
            else:
                print(f"No files found in '{folder_path}'.")
        else:
            print(f"Folder '{folder_path}' does not exist or is not a directory.")
        cat1 = []
        cat1 = file_names
        print(cat1)

        # sum_folder = "summary"
        
        for file_name in os.listdir(folder_path):
            file_path = os.path.join(folder_path,file_name)
            print(file_path)
            
            # Checks if the file is a docx file
            if file_path.endswith('.docx'):
                doc = Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                print(file_name)
                


                
                if file_name in cat1:
                    summ = Chunking.summarize_cat1(Chunking.llm, map_prompt_temp, final_combine_prompt_template, text)

                prompt = '''
                from the following summary extract the company name. just print the company name and nothing else:
                ''' + summ
                file_name_sum = Chunking.openai_api(prompt)
                doc = Document()
                doc.add_paragraph(summ)
                sumpath = os.path.join(sum_folder,file_name_sum)
                doc.save(sumpath + "_Consolidated_Output.docx")
                pp = sumpath + "_Consolidated_Output.docx"
                pp.replace(random_name + "_summary\\",'')
                map = {
                    "file_name" : file_name,
                    "output" : pp
                }
                with open('map.json', 'r') as file:
                    # Load the JSON data
                    data = json.load(file)

                if random_name not in data:
                   data[random_name] = [map] 
                else:
                     data[random_name].append(map)

                with open('map.json', 'w') as file:
                    # Write the modified data back to the JSON file
                    json.dump(data, file, indent=4)