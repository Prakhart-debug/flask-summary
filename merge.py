import os
from docx import Document

class Merging:
    def invoke(folder, base_directory):
        def merge_docx_documents(folder_path, output_file):
            merged_document = Document()

            # List all files in the folder
            files = os.listdir(folder_path)

            # Process each file in the folder
            for filename in files:
                filepath = os.path.join(folder_path, filename)
                if os.path.isfile(filepath) and filename.endswith('.docx'):
                    doc = Document(filepath)
                    # Append all paragraphs from current document to the merged document
                    for paragraph in doc.paragraphs:
                        merged_document.add_paragraph(paragraph.text)

            # Save the merged document to the specified output file
            merged_document.save(output_file)

        #making a folder for combined docx files
        # folder = 'files'
        if not os.path.exists(folder):
            try:
                os.makedirs(folder)
                print(f"Folder '{folder}' successfully created.")
            except Exception as e:
                print(f"Error creating folder '{folder}': {e}")
        else:
            pass

        # base_directory = 'Clusters/'
        for folder_name in os.listdir(base_directory):
            if folder_name.startswith('Cluster') and os.path.isdir(os.path.join(base_directory, folder_name)):
                print(f"Accessing files in '{folder_name}'...")
                
                # Construct the full path to the cluster folder
                cluster_path = os.path.join(base_directory, folder_name)
                output_file = folder + f'/{folder_name}_combined.docx'
                
                merge_docx_documents(cluster_path, output_file)

            else:
                print(f"Skipping '{folder_name}' as it is not a 'Cluster' folder.")
                
        print("DONE")